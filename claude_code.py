import os
import glob
import docx
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.schema import Document

# Set your OpenAI API key
os.environ["OPENAI_API_KEY"] = "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"

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file"""
    doc = docx.Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

def create_resume_pipeline(resumes_folder):
    """Simple pipeline to process resumes and create FAISS vector store"""
    
    # Step 1: Get all DOCX files
    docx_files = glob.glob(f"{resumes_folder}/*.docx")
    print(f"Found {len(docx_files)} resume files")
    
    # Step 2: Extract text and create documents
    documents = []
    for file_path in docx_files:
        filename = os.path.basename(file_path)
        print(f"Processing: {filename}")
        
        # Extract text
        text = extract_text_from_docx(file_path)
        
        # Create document with metadata
        doc = Document(
            page_content=text,
            metadata={"filename": filename, "source": file_path}
        )
        documents.append(doc)
    
    # Step 3: Create embeddings and FAISS vector store
    print("Creating embeddings...")
    embeddings = OpenAIEmbeddings()
    vector_store = FAISS.from_documents(documents, embeddings)
    
    # Step 4: Save the vector store
    vector_store.save_local("resume_vectorstore")
    print("Pipeline complete! Vector store saved.")
    
    return vector_store

def search_resumes(query, k=3):
    """Search for resumes matching the query"""
    embeddings = OpenAIEmbeddings()
    vector_store = FAISS.load_local("resume_vectorstore", embeddings)
    
    results = vector_store.similarity_search(query, k=k)
    
    print(f"\nSearch results for: '{query}'")
    print("-" * 50)
    for i, result in enumerate(results, 1):
        print(f"{i}. {result.metadata['filename']}")
        print(f"   {result.page_content[:200]}...\n")

# Run the pipeline
if __name__ == "__main__":
    # Create the pipeline
    vector_store = create_resume_pipeline("./resumes")  # Change to your folder path
    
    # Example searches
    search_resumes("python developer")
    search_resumes("machine learning engineer")
    search_resumes("project manager")